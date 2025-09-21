import streamlit as st
import hashlib
import os
import PyPDF2
from docx import Document
from dotenv import load_dotenv
from groq import Groq
from serpapi import GoogleSearch
from playwright.async_api import async_playwright
import asyncio
import re
import nltk
import random

# --- Initial Setup ---
# Download NLTK's sentence tokenizer if not already present
try:
    nltk.data.find('tokenizers/punkt')
except nltk.downloader.DownloadError:
    nltk.download('punkt')

# Load environment variables from .env
load_dotenv()

# --- Page Configuration ---
st.set_page_config(
    page_title="Content Source Finder",
    page_icon="🔎",
    layout="wide"
)

st.title("🔎 Content Source Finder & Uniqueness Analyzer")
st.markdown(
    "Upload a file to find out if its sentences exist elsewhere on the web, or upload multiple files to compare them against each other."
)

# --- API Key and Client Initialization ---
groq_api_key = os.getenv("GROQ_API_KEY")
serpapi_api_key = os.getenv("SERPAPI_API_KEY")

if not groq_api_key or not serpapi_api_key:
    missing_keys = [key for key, value in [("GROQ_API_KEY", groq_api_key), ("SERPAPI_API_KEY", serpapi_api_key)] if not value]
    st.error(f"The following API keys are missing in your .env file: {', '.join(missing_keys)}. Please add them.")
    st.stop()

client = Groq(api_key=groq_api_key)

# --- Sidebar and File Uploader ---
with st.sidebar:
    st.header("Upload Document(s)")
    uploaded_files = st.file_uploader(
        "Upload TXT, PDF, or DOCX files.",
        accept_multiple_files=True,
        type=['txt', 'pdf', 'docx']
    )
    analyze_button = st.button("Analyze Content", type="primary", use_container_width=True)

# --- Core Helper Functions ---
def read_file_content(file):
    """Reads content from various file types."""
    try:
        ext = file.name.split('.')[-1].lower()
        if ext == 'txt': return file.getvalue().decode("utf-8")
        if ext == 'docx':
            doc = Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        if ext == 'pdf':
            reader = PyPDF2.PdfReader(file)
            return "".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        st.error(f"Error reading file {file.name}: {e}")
    return ""

def analyze_similarity_with_groq(doc_text, web_text, source_url, model="llama3-70b-8192"):
    """Compares document text against a specific web source."""
    prompt = f"""
    You are a plagiarism analyst. A user's document contains sentences found on the website "{source_url}".
    Analyze the similarity and provide a concise report.

    **Plagiarism Report for Source: {source_url}**
    **Similarity Score**: [An integer from 0 to 100 indicating how similar the user's document is to the content from the source URL.]
    **Analysis Summary**: [A brief, one-sentence summary of your findings.]
    **Copied Sentences Found**:
    *   [List the exact sentence from the user's document that was found on the source website.]
    ---
    **User's Document Snippet:**
    {doc_text[:2000]}
    ---
    **Source Website Content Snippet:**
    {web_text[:2000]}
    ---
    """
    try:
        completion = client.chat.completions.create(messages=[{"role": "user", "content": prompt}], model=model)
        return completion.choices[0].message.content
    except Exception as e:
        st.error(f"Groq API error while analyzing {source_url}: {e}")
    return None

# --- Main Feature Functions ---

def search_web_for_plagiarism(text):
    """
    Splits text into sentences, searches for a sample of them on the web,
    and returns the sentences that were found and their source URLs.
    """
    st.write("Breaking document into sentences for analysis...")
    sentences = nltk.sent_tokenize(text)
    # Filter out very short, non-descriptive sentences
    sentences_to_check = [s for s in sentences if len(s.split()) > 8]

    if not sentences_to_check:
        return []

    # Take a random sample of up to 5 sentences to avoid excessive API calls
    sample_size = min(5, len(sentences_to_check))
    sentences_sample = random.sample(sentences_to_check, sample_size)

    found_sources = []
    progress_bar = st.progress(0, text="Searching web for selected sentences...")

    for i, sentence in enumerate(sentences_sample):
        progress_bar.progress((i + 1) / len(sentences_sample), text=f"Searching for: \"{sentence[:50]}...\"")
        try:
            search = GoogleSearch({
                "q": f'"{sentence}"',
                "api_key": serpapi_api_key
            })
            results = search.get_dict()
            # If organic_results exist, it means we found a match
            if 'organic_results' in results and results['organic_results']:
                first_result = results['organic_results'][0]
                found_sources.append({
                    "sentence": sentence,
                    "url": first_result['link'],
                    "source_title": first_result['title']
                })
        except Exception as e:
            st.warning(f"SerpApi search failed for a sentence. Error: {e}")

    progress_bar.empty()
    return found_sources


async def scrape_text_from_url(url):
    """Uses Playwright to scrape text content from a URL."""
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.goto(url, timeout=15000, wait_until='domcontentloaded')
            content = await page.evaluate("document.body.innerText")
            await browser.close()
            return re.sub(r'\s{2,}', '\n', content).strip()
    except Exception:
        return "" # Silently fail if a page can't be scraped

# --- Application Logic ---

def handle_single_file_analysis(file):
    """The main workflow for checking one file against the web."""
    st.header(f"Web Source Analysis for `{file.name}`")
    content = read_file_content(file)
    if not content:
        st.error("Could not read content from the uploaded file.")
        return

    found_sources = search_web_for_plagiarism(content)

    if not found_sources:
        st.success("✅ **Content Appears to be Unique**")
        st.balloons()
        st.write("After checking several key sentences from your document, no direct matches were found on the web.")
        return

    st.error("🚨 **Potential Plagiarism Detected!**")
    st.write("The following sentences from your document were found on other websites. This indicates that the content may not be original.")

    # Display the table of found sentences and their sources
    st.subheader("Found Sentences and Their Online Sources")
    for source in found_sources:
        with st.container(border=True):
            st.markdown(f"**Sentence:** *\"{source['sentence']}\"*")
            st.markdown(f"**Found At:** [{source['source_title']}]({source['url']})")

    # Perform a deep-dive analysis on the first source found
    st.subheader("In-Depth Analysis")
    first_source = found_sources[0]
    with st.spinner(f"Performing deep analysis against the first source: {first_source['url']}"):
        web_content = asyncio.run(scrape_text_from_url(first_source['url']))
        if web_content:
            report = analyze_similarity_with_groq(content, web_content, first_source['url'])
            if report:
                st.info(report)
        else:
            st.warning(f"Could not scrape content from {first_source['url']} to perform a detailed comparison.")


def handle_multi_file_analysis(files):
    """The original workflow for comparing multiple files against each other."""
    st.header("Pairwise Document Comparison Report")
    # This function remains the same as in the previous version.
    # It reads multiple files and compares each pair.
    # [Code for this function is omitted for brevity but is unchanged from your last version]
    st.info("Multi-file comparison logic is ready.")


# --- Main Execution Block ---
if analyze_button and uploaded_files:
    if len(uploaded_files) == 1:
        handle_single_file_analysis(uploaded_files[0])
    elif len(uploaded_files) >= 2:
        # For this example, we're focusing on the single-file logic.
        # You can paste your existing multi-file handler function here.
        st.warning("Multi-file comparison is not shown in this example. Please use the single-file checker.")
        # handle_multi_file_analysis(uploaded_files) # Uncomment this to enable it
else:
    st.info("Upload a document and click 'Analyze Content' to find its sources on the web.")
