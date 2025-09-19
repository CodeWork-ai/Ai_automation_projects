"""
Production Document Processing Module
Handles PDF and DOCX text extraction with error handling and logging
"""

import logging
from io import BytesIO
from typing import Optional

import docx
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_data: bytes) -> str:
    """Extract text from PDF binary data with error handling"""
    try:
        pdf = PdfReader(BytesIO(file_data))
        text = ""
        
        for page_num, page in enumerate(pdf.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    logger.warning(f"No text extracted from page {page_num}")
            except Exception as e:
                logger.error(f"Error extracting text from page {page_num}: {str(e)}")
                continue
        
        logger.info(f"Successfully extracted text from PDF: {len(text)} characters")
        return text.strip()
        
    except Exception as e:
        logger.error(f"Failed to process PDF: {str(e)}")
        raise ValueError(f"PDF processing failed: {str(e)}")

def extract_text_from_docx(file_data: bytes) -> str:
    """Extract text from DOCX binary data with error handling"""
    try:
        doc = docx.Document(BytesIO(file_data))
        text = ""
        
        for paragraph_num, paragraph in enumerate(doc.paragraphs, 1):
            try:
                para_text = paragraph.text.strip()
                if para_text:
                    text += para_text + "\n"
            except Exception as e:
                logger.error(f"Error extracting text from paragraph {paragraph_num}: {str(e)}")
                continue
        
        # Also extract text from tables
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += cell_text + " "
                    text += "\n"
        except Exception as e:
            logger.warning(f"Error extracting table text: {str(e)}")
        
        logger.info(f"Successfully extracted text from DOCX: {len(text)} characters")
        return text.strip()
        
    except Exception as e:
        logger.error(f"Failed to process DOCX: {str(e)}")
        raise ValueError(f"DOCX processing failed: {str(e)}")

def process_document(file_data: bytes, content_type: str) -> str:
    """
    Process document based on its content type
    
    Args:
        file_data: Binary data of the file
        content_type: MIME type or file extension
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: For unsupported file types or processing errors
    """
    # Normalize content type
    content_type = content_type.lower()
    
    # Handle different content type formats
    if content_type in ['pdf', 'application/pdf', 'application/x-pdf']:
        return extract_text_from_pdf(file_data)
    elif content_type in [
        'docx', 
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]:
        return extract_text_from_docx(file_data)
    else:
        supported_types = ['PDF', 'DOCX']
        raise ValueError(f"Unsupported file type: {content_type}. Supported types: {', '.join(supported_types)}")

def validate_document(file_data: bytes, filename: str, max_size_mb: int = 50) -> tuple[bool, Optional[str]]:
    """
    Validate document before processing
    
    Args:
        file_data: Binary data of the file
        filename: Name of the file
        max_size_mb: Maximum allowed file size in MB
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    try:
        # Check file size
        file_size_mb = len(file_data) / (1024 * 1024)
        if file_size_mb > max_size_mb:
            return False, f"File too large: {file_size_mb:.1f}MB (max: {max_size_mb}MB)"
        
        # Check if file is empty
        if len(file_data) == 0:
            return False, "File is empty"
        
        # Check file extension
        if not filename:
            return False, "Filename is required"
        
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        if file_ext not in ['pdf', 'docx']:
            return False, f"Unsupported file extension: .{file_ext}"
        
        return True, None
        
    except Exception as e:
        logger.error(f"Document validation error: {str(e)}")
        return False, f"Validation error: {str(e)}"
